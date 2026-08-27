"""Parse DOCX timetable tables into the shared timetable row format."""

from __future__ import annotations

import pandas as pd
from docx import Document

from apps.timetable.services.excel_parser import normalise_dataframe


def parse_docx(file_path: str) -> list[dict]:
    """Parse the largest table in a DOCX timetable document."""
    document = Document(file_path)
    if not document.tables:
        raise ValueError("No tables found in the DOCX timetable.")

    table = max(document.tables, key=lambda candidate: len(candidate.rows))
    table_rows = [
        [cell.text.strip() for cell in row.cells]
        for row in table.rows
    ]
    if len(table_rows) <= 1:
        raise ValueError("The DOCX timetable table has no data rows.")

    headers = table_rows[0]
    data = table_rows[1:]
    return normalise_dataframe(pd.DataFrame(data, columns=headers))